# Ai_Premium_Content_Generator.py
# All-in-One Premium AI Content Generator for Pyroid (Android)

import os
import time
from dotenv import load_dotenv
import streamlit as st
from pyngrok import ngrok
from openai import OpenAI
from docx import Document
import plotly.express as px
from googletrans import Translator

# -------------------------
# Load environment variables
# -------------------------
load_dotenv()

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
NGROK_AUTH_TOKEN = os.getenv("NGROK_AUTH_TOKEN")

if not OPENAI_API_KEY or not NGROK_AUTH_TOKEN:
    st.error("Please set OPENAI_API_KEY and NGROK_AUTH_TOKEN in .env file.")
    st.stop()

# -------------------------
# Initialize OpenAI client
# -------------------------
client = OpenAI(api_key=OPENAI_API_KEY)

# -------------------------
# Start ngrok for Streamlit
# -------------------------
ngrok.set_auth_token(NGROK_AUTH_TOKEN)
public_url = ngrok.connect(8501)
st.sidebar.success(f"🌐 Public URL: {public_url}")

# -------------------------
# Streamlit App
# -------------------------
st.set_page_config(
    page_title="All-in-One Premium AI Content Generator",
    page_icon="🌟",
    layout="wide"
)

st.title("🌟 All-in-One Premium AI Content Generator")

# --- Dashboard Metrics ---
if "metrics" not in st.session_state:
    st.session_state["metrics"] = {"generated": 0, "categories": {}, "recent_topics": []}

# --- Settings Panel ---
st.sidebar.header("⚙️ Settings")
language = st.sidebar.selectbox("Language", ["English", "Spanish", "French", "German", "Chinese", "Japanese", "Hindi"])
content_type = st.sidebar.selectbox("Content Type", ["Blog Post", "Article", "Social Media Post"])
num_variations = st.sidebar.slider("Number of Variations", 1, 5, 1)
word_count = st.sidebar.slider("Word Count", 50, 1000, 200)
category = st.sidebar.text_input("Category", "General")
tone = st.sidebar.selectbox("Tone / Style", ["Informative", "Casual", "Professional", "Funny"])
seo_keywords = st.sidebar.text_input("SEO Keywords (comma separated)")

# --- Content Input ---
st.subheader("Enter Topic / Prompt")
topic = st.text_input("Topic / Prompt", "")

image_prompt = st.text_input("Optional Image Prompt for AI Image Generation", "")

if st.button("Generate Content") and topic:
    with st.spinner("Generating content..."):
        results = []
        translator = Translator()
        for i in range(num_variations):
            prompt_text = f"Generate a {content_type} in {language} on '{topic}' with a {tone} tone, approx {word_count} words, under {category} category."
            if seo_keywords:
                prompt_text += f" Include these SEO keywords: {seo_keywords}."

            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt_text}]
            )
            content = response.choices[0].message.content
            results.append(content)
            
            # Update metrics
            st.session_state.metrics["generated"] += 1
            st.session_state.metrics["recent_topics"].append(topic)
            st.session_state.metrics["categories"][category] = st.session_state.metrics["categories"].get(category, 0) + 1

        # Display results
        for idx, text in enumerate(results):
            st.markdown(f"### Variation {idx+1}")
            st.write(text)

        # Optional: Export to DOCX
        doc = Document()
        doc.add_heading(f"{topic} - Generated Content", 0)
        for idx, text in enumerate(results):
            doc.add_heading(f"Variation {idx+1}", level=1)
            doc.add_paragraph(text)
        doc_file = f"{topic.replace(' ', '_')}_content.docx"
        doc.save(doc_file)
        st.success(f"✅ DOCX saved as {doc_file}")

# --- Dashboard ---
st.sidebar.header("📊 Dashboard")
st.sidebar.write(f"Total contents generated: {st.session_state.metrics['generated']}")
st.sidebar.write("Recent Topics")
for t in st.session_state.metrics["recent_topics"][-5:]:
    st.sidebar.write(f"• {t}")

# --- Plot category metrics ---
if st.session_state.metrics["categories"]:
    fig = px.bar(
        x=list(st.session_state.metrics["categories"].keys()),
        y=list(st.session_state.metrics["categories"].values()),
        labels={"x":"Category", "y":"Generated Count"},
        title="Generated Content by Category"
    )
    st.plotly_chart(fig)